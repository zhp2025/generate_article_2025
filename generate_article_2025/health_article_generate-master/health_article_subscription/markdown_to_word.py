import os
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from bs4 import BeautifulSoup
from bs4.element import Tag # 修正: 直接从 bs4.element 导入 Tag
from pathlib import Path
import threading
import queue
import logging

# --- 配置常量 (Configuration Constants) ---
DEFAULT_FONT_NAME = 'Arial'
DEFAULT_FONT_SIZE = Pt(11)
H1_FONT_SIZE = Pt(11)
H2_FONT_SIZE = Pt(11)
TITLE_FONT_SIZE = Pt(14)
CODE_FONT_NAME = 'Courier New'
CODE_FONT_SIZE = Pt(10)
FIRST_LINE_INDENT_PT = Pt(21)
DISCLAIMER_FONT_SIZE = Pt(10)

# --- 设置日志 (Setup Logging) ---
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
# 创建控制台处理器 (console handler)
console_handler = logging.StreamHandler()
console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
logger.addHandler(console_handler)
# 如果需要，也可以添加文件处理器 (file handler)
# file_handler = logging.FileHandler('conversion.log', encoding='utf-8')
# file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
# logger.addHandler(file_handler)


class MarkdownToWordConverter:
    def __init__(self, source_dir: Path, output_dir: Path):
        self.source_dir = source_dir
        self.output_dir = output_dir

    def clean_text(self, text: str) -> str:
        """
        清理文本，移除与XML或Word处理不兼容的字符，同时保留Unicode字符（如中文）。
        """
        if not isinstance(text, str):
            return ''

        # 1. 移除C0控制字符（除了制表符\t, 换行符\n, 回车符\r）
        #    特别是空字符 \x00，它对于XML和python-docx来说是无效的。
        #    \x7f (DEL) 也一并移除。
        cleaned_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        # 2. 移除Unicode中的一些特殊格式字符，如零宽空格、字节顺序标记(BOM)等。
        cleaned_text = re.sub(r'[\u200b-\u200f\u202a-\u202e\uFEFF]', '', cleaned_text)

        return cleaned_text

    def preserve_references(self, md_content: str) -> str:
        """
        将参考文献部分包裹在三反引号中，以便在Markdown转HTML时保留其原始格式。
        """
        pattern = r'(参考文献[:：]\s*\n(?:\[[0-9]+\].*\n?)+)'
        def replacer(match: re.Match) -> str:
            return f"```\n{match.group(1).strip()}\n```"
        return re.sub(pattern, replacer, md_content)

    def process_html_element(self, element: Tag, doc: Document):
        """处理HTML元素并将其添加到Word文档中。"""
        if element.name == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(element.get_text(strip=True))
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.size = H1_FONT_SIZE
        elif element.name == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(element.get_text(strip=True))
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = H2_FONT_SIZE
        elif element.name == 'h3':
            doc.add_heading(element.get_text(strip=True), level=3)
        elif element.name == 'p':
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT_PT
            for child in element.contents:
                if isinstance(child, str):
                    if child.strip():
                        p.add_run(child.strip())
                elif child.name in ['strong', 'b']:
                    p.add_run(child.get_text(strip=True)).bold = True
                elif child.name in ['em', 'i']:
                    p.add_run(child.get_text(strip=True)).italic = True
                elif child.name == 'code':
                    run = p.add_run(child.get_text(strip=True))
                    run.font.name = CODE_FONT_NAME
                    run.font.size = CODE_FONT_SIZE
                elif child.name == 'a':
                    link_text = child.get_text(strip=True)
                    link_href = child.get('href', '')
                    p.add_run(f"{link_text} ({link_href})")
        elif element.name == 'pre':
            code_block_text = element.get_text()
            if "参考文献" in code_block_text:
                for line in code_block_text.strip().split('\n'):
                    p_ref = doc.add_paragraph(line)
            else:
                p = doc.add_paragraph()
                run = p.add_run(code_block_text)
                run.font.name = CODE_FONT_NAME
                run.font.size = CODE_FONT_SIZE
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                list_item_text = li.get_text(separator=' ', strip=True)
                doc.add_paragraph(
                    list_item_text,
                    style='ListBullet' if element.name == 'ul' else 'ListNumber'
                )
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: return
            num_cols = max(len(row.find_all(['td', 'th'])) for row in rows) if rows else 0
            if num_cols == 0:
                logger.warning("Skipping table with no columns found.")
                return
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            for i, row_element in enumerate(rows):
                cells = row_element.find_all(['td', 'th'])
                for j, cell_element in enumerate(cells):
                    if j < num_cols:
                        table.cell(i, j).text = cell_element.get_text(strip=True)

    def convert_md_to_word(self, md_file_path: Path, output_file_path: Path) -> bool:
        """将单个Markdown文件转换为Word文件。"""
        try:
            md_content = md_file_path.read_text(encoding='utf-8')
            md_content = self.clean_text(md_content)
            md_content = self.preserve_references(md_content)

            html_content = markdown.markdown(
                md_content,
                extensions=['tables', 'fenced_code', 'nl2br']
            )

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = DEFAULT_FONT_NAME
            style.font.size = DEFAULT_FONT_SIZE

            file_title = md_file_path.stem
            title_paragraph = doc.add_paragraph()
            title_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_paragraph.add_run(file_title)
            run.bold = True
            run.font.size = TITLE_FONT_SIZE

            soup = BeautifulSoup(html_content, 'html.parser')
            for element in soup.children:
                if isinstance(element, Tag):
                    self.process_html_element(element, doc)
                elif isinstance(element, str) and element.strip():
                     doc.add_paragraph(element.strip())

            disclaimer = "（注：本文是根据权威医学资料结合个人观点撰写的内容，人名均为化名，部分图片为网图，文章禁止转载、抄袭。文中所有建议需遵从医嘱）"
            p_disclaimer = doc.add_paragraph(disclaimer)
            p_disclaimer.paragraph_format.space_before = Pt(12)
            p_disclaimer.paragraph_format.space_after = Pt(12)
            p_disclaimer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p_disclaimer.runs:
                p_disclaimer.runs[0].italic = True
                p_disclaimer.runs[0].font.size = DISCLAIMER_FONT_SIZE
            
            output_file_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_file_path)
            return True
        except Exception as e:
            logger.error(f"转换文件 {md_file_path} 时出错: {e}", exc_info=True)
            return False

    def process_directory(self, progress_queue: queue.Queue = None) -> tuple[int, int]:
        """遍历源目录，转换所有Markdown文件。"""
        success_count, error_count = 0, 0
        md_files = list(self.source_dir.rglob('*.md'))
        total_files = len(md_files)

        for i, md_file_path in enumerate(md_files):
            rel_path = md_file_path.relative_to(self.source_dir).parent
            output_root = self.output_dir / rel_path
            output_file_name = md_file_path.stem + '.docx'
            output_file_path = output_root / output_file_name

            if progress_queue:
                progress_queue.put(f"处理中: {md_file_path.name} ({i+1}/{total_files})")
            
            logger.info(f"正在转换 {md_file_path} -> {output_file_path}")

            if self.convert_md_to_word(md_file_path, output_file_path):
                success_count += 1
            else:
                error_count += 1
        
        summary_success = f"成功转换: {success_count} 个文件"
        summary_error = f"转换失败: {error_count} 个文件"
        logger.info("\n转换完成:")
        logger.info(summary_success)
        logger.info(summary_error)
        
        if progress_queue:
            progress_queue.put("DONE")
            progress_queue.put(f"转换完成！\n{summary_success}\n{summary_error}")

        return success_count, error_count


class ConverterGUI:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Markdown to Word Converter")
        self.window.geometry("700x500")
        self.converter_thread = None
        self.progress_queue = queue.Queue()
        self.create_widgets()
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)

    def create_widgets(self):
        main_frame = tk.Frame(self.window, padx=10, pady=10)
        main_frame.pack(fill=tk.BOTH, expand=True)

        source_frame = tk.Frame(main_frame)
        source_frame.pack(pady=5, fill=tk.X)
        tk.Label(source_frame, text="源文件夹:", width=15, anchor="w").pack(side=tk.LEFT)
        self.source_path_var = tk.StringVar()
        tk.Entry(source_frame, textvariable=self.source_path_var).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        tk.Button(source_frame, text="选择...", command=self.select_source).pack(side=tk.LEFT)

        target_frame = tk.Frame(main_frame)
        target_frame.pack(pady=5, fill=tk.X)
        tk.Label(target_frame, text="目标文件夹:", width=15, anchor="w").pack(side=tk.LEFT)
        self.target_path_var = tk.StringVar()
        tk.Entry(target_frame, textvariable=self.target_path_var).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        tk.Button(target_frame, text="选择...", command=self.select_target).pack(side=tk.LEFT)

        self.convert_button = tk.Button(main_frame, text="开始转换", command=self.start_conversion_thread, width=20, height=2)
        self.convert_button.pack(pady=15)

        tk.Label(main_frame, text="日志/状态:", anchor="w").pack(pady=(5,0), fill=tk.X)
        self.log_text = scrolledtext.ScrolledText(main_frame, height=10, wrap=tk.WORD, state=tk.DISABLED)
        self.log_text.pack(pady=5, fill=tk.BOTH, expand=True)

    def _log_message(self, message: str, clear: bool = False):
        self.log_text.config(state=tk.NORMAL)
        if clear:
            self.log_text.delete(1.0, tk.END)
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.config(state=tk.DISABLED)
        self.window.update_idletasks()

    def select_source(self):
        folder = filedialog.askdirectory(title="选择源文件夹")
        if folder:
            self.source_path_var.set(folder)
            self._log_message(f"已选择源文件夹: {folder}")

    def select_target(self):
        folder = filedialog.askdirectory(title="选择目标文件夹")
        if folder:
            self.target_path_var.set(folder)
            self._log_message(f"已选择目标文件夹: {folder}")

    def check_queue(self):
        try:
            while True:
                message = self.progress_queue.get_nowait()
                if message == "DONE":
                    self.conversion_finished(success=True)
                    return
                elif isinstance(message, tuple) and message[0] == "ERROR":
                    self.conversion_finished(success=False, error_message=message[1])
                    return
                else:
                    self._log_message(message)
        except queue.Empty:
            pass
        
        if self.converter_thread and self.converter_thread.is_alive():
            self.window.after(100, self.check_queue)
        elif self.converter_thread and self.convert_button['state'] == tk.DISABLED:
            self.conversion_finished(success=False, error_message="转换进程意外终止。")

    def start_conversion_thread(self):
        source_str, target_str = self.source_path_var.get(), self.target_path_var.get()
        if not source_str or not target_str:
            messagebox.showerror("错误", "请选择源文件夹和目标文件夹。")
            return
        
        source_path, target_path = Path(source_str), Path(target_str)
        if not source_path.is_dir():
            messagebox.showerror("错误", f"源路径不是一个有效的文件夹:\n{source_path}")
            return

        self._log_message("开始转换...", clear=True)
        self.convert_button.config(state=tk.DISABLED)
        
        while not self.progress_queue.empty():
            try: self.progress_queue.get_nowait()
            except queue.Empty: break

        converter = MarkdownToWordConverter(source_path, target_path)
        self.converter_thread = threading.Thread(
            target=self.run_conversion_task, args=(converter,), daemon=True
        )
        self.converter_thread.start()
        self.window.after(100, self.check_queue)

    def run_conversion_task(self, converter: MarkdownToWordConverter):
        """线程执行的实际转换任务。"""
        try:
            converter.process_directory(self.progress_queue)
        except Exception as e:
            logger.error(f"转换线程中出现未处理的错误: {e}", exc_info=True)
            self.progress_queue.put(("ERROR", f"发生意外错误: {e}"))

    def conversion_finished(self, success: bool, error_message: str = None):
        self.convert_button.config(state=tk.NORMAL)
        self.converter_thread = None
        
        if success:
            messagebox.showinfo("完成", "文档转换已完成，详情请查看日志。")
        else:
            final_msg = "转换过程中出现错误。"
            if error_message:
                final_msg += f"\n详情: {error_message}"
            self._log_message(final_msg)
            messagebox.showerror("错误", f"{final_msg}\n请检查日志或控制台输出。")
            
    def on_closing(self):
        """处理窗口关闭事件。"""
        if self.converter_thread and self.converter_thread.is_alive():
            if messagebox.askokcancel("退出", "转换正在进行中，您确定要退出吗？"):
                self.window.destroy()
        else:
            self.window.destroy()

    def run(self):
        self.window.mainloop()

def main():
    logger.info("应用程序启动。")
    gui = ConverterGUI()
    gui.run()
    logger.info("应用程序关闭。")

if __name__ == '__main__':
    main()