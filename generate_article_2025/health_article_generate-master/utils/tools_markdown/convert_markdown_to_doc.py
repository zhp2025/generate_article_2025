import os
from tqdm import tqdm
import re
from docx import Document
from markdown import markdown
from docx.shared import Pt
from pathlib import Path
from bs4 import BeautifulSoup

# 全局变量初始化
list_level = 0  # 添加全局变量初始化

def convert_markdown_to_word(input_dir, output_dir):
    global list_level  # 声明使用全局变量
    list_level = 0  # 每次转换前重置层级
    
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    for filename in tqdm(os.listdir(input_dir)):
        if filename.lower().endswith('.md'):
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, f"{Path(filename).stem}.docx")
            
            try:
                with open(input_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                
                soup = BeautifulSoup(markdown(md_content), 'html.parser')
                doc = Document()
                set_default_font(doc)
                
                # 重置当前文件的列表层级
                list_level = 0
                
                process_element(soup, doc)
                doc.save(output_path)
                print(f"转换成功：{filename}")
                
            except Exception as e:
                print(f"转换失败 {filename}: {str(e)}")

def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)

def process_element(element, doc):
    global list_level  # 声明全局变量
    
    if element.name == 'h1':
        doc.add_heading(element.text.strip(), level=0)
    elif element.name == 'h2':
        doc.add_heading(element.text.strip(), level=1)
    elif element.name == 'p':
        process_paragraph(element, doc)
    elif element.name in ['ul', 'ol']:
        # 进入列表时增加层级
        original_level = list_level
        list_level += 1
        for child in element.children:
            process_element(child, doc)
        # 恢复层级
        list_level = original_level
    elif element.name == 'li':
        process_list_item(element, doc)
    # 处理嵌套元素
    if hasattr(element, 'children'):
        for child in element.children:
            if child != '\n':
                process_element(child, doc)

def process_paragraph(element, doc):
    p = doc.add_paragraph()
    for content in element.contents:
        if content.name == 'strong':
            add_bold_text(p, content.text)
        else:
            add_normal_text(p, str(content))

def process_list_item(element, doc):
    p = doc.add_paragraph()
    # 根据层级设置缩进
    p.paragraph_format.left_indent = Pt(20 * list_level)
    
    # 判断列表类型
    if element.find_parent('ul'):
        p.style = 'List Bullet'
    elif element.find_parent('ol'):
        p.style = 'List Number'
    
    for content in element.contents:
        if content.name == 'strong':
            add_bold_text(p, content.text)
        else:
            add_normal_text(p, str(content))

def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True

def add_normal_text(paragraph, text):
    clean_text = re.sub(r'<[^>]+>', '', text)
    paragraph.add_run(clean_text)


if __name__ == "__main__":
    input_dir = "/root/project/product_promo_article_generate/health_article_generate/util_test/extension_test/examples/output-2月28"
    output_dir = "word_documents"
    convert_markdown_to_word(input_dir, output_dir)